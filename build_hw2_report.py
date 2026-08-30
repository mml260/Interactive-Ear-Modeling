from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).parent / 'docs' / 'HW2_Middle_Ear_Report_Draft.docx'

BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
INK = '0B2545'
MUTED = '5B6F82'
TABLE_FILL = 'F2F4F7'
CALLOUT_FILL = 'F4F6F9'
PLACEHOLDER_FILL = 'F8FAFC'
BORDER = 'B7C8D9'


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in('w:tcMar')
    if mar is None:
        mar = OxmlElement('w:tcMar')
        tc_pr.append(mar)
    for edge, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = mar.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        table_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths_dxa)))
    tbl_w.set(qn('w:type'), 'dxa')
    ind = table_pr.first_child_found_in('w:tblInd')
    if ind is None:
        ind = OxmlElement('w:tblInd')
        table_pr.append(ind)
    ind.set(qn('w:w'), str(indent))
    ind.set(qn('w:type'), 'dxa')
    layout = table_pr.first_child_found_in('w:tblLayout')
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        table_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    grid = table._tbl.tblGrid
    for column, width in zip(grid.gridCol_lst, widths_dxa):
        column.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)


def set_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = borders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), color)


def add_para(doc, text='', style=None, size=None, color=INK, bold=False, italic=False, align=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_inline_label(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    label_run = p.add_run(label)
    set_run_font(label_run, size=11, color=INK, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=11, color=INK)
    return p


def add_figure_placeholder(doc, title, capture):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_borders(table)
    cell = table.cell(0, 0)
    shade_cell(cell, PLACEHOLDER_FILL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f'[ INSERT SCREENSHOT: {title} ]')
    set_run_font(run, size=10, color=DARK_BLUE, bold=True)
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption_run = caption.add_run(f'{title}. ')
    set_run_font(caption_run, size=9, color=INK, bold=True)
    capture_run = caption.add_run(capture)
    set_run_font(capture_run, size=9, color=MUTED, italic=True)


def add_experiment(doc, title, change, prediction, result, conclusion, figure_title, figure_capture):
    add_para(doc, title, style='Heading 2')
    add_inline_label(doc, 'Controlled change: ', change)
    add_inline_label(doc, 'Pre-run prediction: ', prediction)
    add_inline_label(doc, 'Calculated result: ', result)
    add_inline_label(doc, 'Comparison: ', conclusion)
    add_figure_placeholder(doc, figure_title, figure_capture)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run, size=8.5, color=MUTED)


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ('Heading 1', 16, BLUE, 16, 8),
        ('Heading 2', 13, BLUE, 12, 6),
        ('Heading 3', 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run('EAR ACOUSTICS SIMULATOR  |  HOMEWORK 2')
    set_run_font(header_run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run('Middle-Ear Extension Report  |  Page ')
    set_run_font(footer_run, size=8.5, color=MUTED)
    add_page_field(footer)

    add_para(doc, 'HOMEWORK 2', size=10, color=BLUE, bold=True, after=3)
    add_para(doc, 'Middle-Ear Extension Report', size=24, color=INK, bold=True, after=5)
    add_para(doc, 'Ear Acoustics Simulator', size=13, color=MUTED, italic=True, after=14)

    metadata = doc.add_table(rows=3, cols=2)
    set_table_geometry(metadata, [1800, 7560])
    set_borders(metadata)
    for row, label, value in zip(metadata.rows, ['Student', 'Course', 'Date'], ['[Your name]', '[Course name]', '[Date]']):
        shade_cell(row.cells[0], TABLE_FILL)
        row.cells[0].text = ''
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=10, color=DARK_BLUE, bold=True)
        row.cells[1].text = ''
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=10, color=INK)

    add_para(doc, '1. Objective and Model Scope', style='Heading 1')
    add_para(doc, 'This project extends the existing outer-ear simulator with a standalone TypeScript middle-ear model. The modeled signal path is pinna/concha → ear canal → eardrum → ossicles → cochlear load. The app can display and sonify the outer-ear response, the middle-ear response, or their cascade.')
    add_inline_label(doc, 'Cascade assumption: ', 'The models are connected as H_total(f) = H_outer(f) × H_middle(f). This is an explicit approximation because loading between independently modeled stages is not solved.')
    add_figure_placeholder(doc, 'Figure 1 — Outer, middle, and combined magnitude responses', 'Take three app screenshots at 20 Hz–20 kHz: Listening path = Outer ear, Middle ear (OME = 0), and Both modules. Place them side-by-side or vertically here.')

    add_para(doc, '2. Middle-Ear Circuit and Verification', style='Heading 1')
    add_para(doc, 'The course equivalent circuit was implemented as separate functions for Z_cavity, Z_eardrum, Z_ossicles, Z_joint, and Z_cochlea. The output transfer is the complex pressure ratio from the model input to the cochlear-load node.')
    equations = doc.add_table(rows=1, cols=1)
    set_table_geometry(equations, [9360])
    set_borders(equations)
    cell = equations.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    cell.text = ''
    eq = cell.paragraphs[0]
    eq.paragraph_format.space_after = Pt(0)
    eq_text = ('Z_cavity = (Z_Cp + Z_La + Z_Ra) || Z_Rm || Z_Ct\n'
               'Z_eardrum = Z_CD1 + (Z_LD || (Z_CD2 + Z_RD2)) + Z_RD1\n'
               'Z_ossicles = Z_CO + Z_LO + Z_RO;  Z_load = Z_joint || Z_cochlea\n'
               'Z_middle = Z_cavity + [Z_eardrum || (Z_ossicles + Z_load)]\n'
               'H_middle = (Z_after_A / Z_middle) × (Z_load / (Z_ossicles + Z_load))')
    eq_run = eq.add_run(eq_text)
    set_run_font(eq_run, size=9.5, color=INK)
    add_para(doc, 'Topology verification was performed in code rather than by comparing screenshots. At 1 kHz, tests independently reconstruct every complex branch from the impedance equations, then compare the real and imaginary components of each subsystem, Z_middle, and H_middle against the implementation. Primitive tests also verify Z_R = R, Z_L = jωL, and Z_C = 1/(jωC). The completed project has 21 automated tests passing, plus a successful production build and lint check.')

    add_para(doc, '3. Component Values and Classifications', style='Heading 1')
    add_para(doc, 'Exact instructor values were not available, so the initial parameter set is a source-tagged, replaceable Zwislocki-style historical analogue baseline. It produces a broad modeled peak near 0.87 kHz, close to the course figure’s requested approximately 1 kHz behavior. These are not claimed as individual anatomical measurements.')
    classification = doc.add_table(rows=1, cols=3)
    set_table_geometry(classification, [2400, 2400, 4560])
    set_borders(classification)
    headers = ['Classification', 'Example', 'Reason']
    for cell, text in zip(classification.rows[0].cells, headers):
        shade_cell(cell, TABLE_FILL)
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
    for category, example, reason in [
        ('Physically derived', 'Cp = 5.1 μF; Ct = 0.35 μF', 'Cavity-compliance terms tied to estimated volume behavior in the analogue.'),
        ('Fitted historical analogue', 'LO = 40 mH; RO = 70 Ω; CD1 = 0.23 μF', 'Historical aggregate values used to approximate middle-ear impedance behavior.'),
        ('Assumed simplification', 'LC = 0 mH', 'Cochlear-load inertance is omitted rather than invented; this is an explicit model assumption.'),
    ]:
        row = classification.add_row()
        for cell, text in zip(row.cells, [category, example, reason]):
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            set_run_font(run, size=9.2, color=INK)

    add_para(doc, '4. AI-Assisted Parameter Experiments', style='Heading 1')
    add_para(doc, 'Each formal case changes one parameter while resetting the other two experiment controls and the OME scenario to baseline. The predictions below were made before running the model, then compared with the calculated response.')
    add_experiment(
        doc,
        'Experiment 1 — Ossicular inertance (mass)',
        'LO: 40 → 60 mH (+50%).',
        'More inertance should shift mass-related behavior and reduce high-frequency transmission relative to the broad middle-frequency peak.',
        'Peak frequency moved 871 → 752 Hz. Relative to baseline, the response changed −1.68 dB at 2 kHz and −2.95 dB at 4 kHz.',
        'The expected downward frequency shift and high-frequency reduction occurred. The small peak-height increase is a network interaction, not global amplification.',
        'Figure 2 — Inertance comparison',
        'Capture baseline first, then click Load case for Experiment 1. Insert the two magnitude plots as a side-by-side comparison.',
    )
    doc.add_page_break()
    add_experiment(
        doc,
        'Experiment 2 — Eardrum compliance',
        'CD1: 0.23 → 0.50 μF (+117%).',
        'More eardrum compliance should change the eardrum-branch loading and alter transfer near the broad peak.',
        'Peak frequency moved 871 → 732 Hz; peak magnitude changed −1.49 → −1.82 dB. The response at 1 kHz was −0.73 dB relative to baseline.',
        'The response moved downward and decreased around 1 kHz, consistent with the prediction.',
        'Figure 3 — Compliance comparison',
        'Capture baseline first, then click Load case for Experiment 2. Insert the two magnitude plots as a side-by-side comparison.',
    )
    add_experiment(
        doc,
        'Experiment 3 — Ossicular loss (resistance)',
        'RO: 70 → 300 Ω (+329%).',
        'More ossicular loss should dissipate more energy and reduce transmission near the broad resonance.',
        'Peak frequency moved 871 → 414 Hz; peak magnitude changed −1.49 → −3.83 dB. The response at 1 kHz was −2.69 dB relative to baseline.',
        'The predicted reduction occurred; this was the strongest loss-control effect in the selected topology.',
        'Figure 4 — Loss comparison',
        'Capture baseline first, then click Load case for Experiment 3. Insert the two magnitude plots as a side-by-side comparison.',
    )

    doc.add_page_break()
    add_para(doc, '5. Graduate Extension: OME-Like Transfer Scenario', style='Heading 1')
    add_para(doc, 'The app includes a small-child otitis media with effusion (OME) sensitivity scenario. Its continuous severity control reduces eardrum and joint compliance while increasing ossicular inertance and loss. When severity is above zero, the app overlays the baseline and OME-like responses and applies the same profile to playback.')
    add_para(doc, 'This is a mechanism-based educational sensitivity model, not a pediatric normative fit, patient simulation, or diagnostic tool. The endpoint multipliers are transparent and can be replaced if course-specific values become available.')
    add_figure_placeholder(doc, 'Figure 5 — Small-child OME-like transfer comparison', 'Select Middle ear, then raise OME severity to 1.0. Capture the baseline-versus-OME overlay; include the severity value in the caption.')

    doc.add_page_break()
    add_para(doc, '6. AI Use and Engineering Log', style='Heading 1')
    log = doc.add_table(rows=1, cols=3)
    set_table_geometry(log, [2000, 3600, 3760])
    set_borders(log)
    for cell, text in zip(log.rows[0].cells, ['Question / AI use', 'Verification', 'Decision']):
        shade_cell(cell, TABLE_FILL)
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
    records = [
        ('Circuit interpretation', 'Series/parallel relations were reconstructed from the supplied course circuit and checked with independent complex-impedance tests at 1 kHz.', 'Used the verified topology; did not treat an AI interpretation as sufficient evidence.'),
        ('Initial component values', 'Exact course values were unavailable. Historical analogue values were source-tagged, labeled provisional, and made replaceable in the app.', 'Retained a transparent Zwislocki-style baseline rather than presenting invented values as physiology.'),
        ('Audio FFT suggestion', 'The first custom FFT filtering approach failed the audio-processing verification tests.', 'Rejected that implementation and replaced it with the tested fft.js library. This is the required real AI correction/rejection example.'),
    ]
    for record in records:
        row = log.add_row()
        for cell, text in zip(row.cells, record):
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            set_run_font(run, size=9.0, color=INK)

    add_para(doc, '7. Limitations', style='Heading 1')
    add_para(doc, 'The complete outer + middle-ear model has the following important limitations:')
    for limitation in [
        'It is a lumped, small-signal equivalent circuit, not a full anatomical or individual-ear model.',
        'The middle-ear values are a historical aggregate analogue baseline and have limited high-frequency fidelity; they are not measured for this project or for a specific listener.',
        'The outer and middle responses are cascaded independently, so loading between those stages is not solved.',
        'The OME scenario is a transparent educational sensitivity profile, not a clinical calibration or pediatric diagnostic model.',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(limitation)
        set_run_font(r, size=11, color=INK)

    add_para(doc, 'References', style='Heading 1')
    for source in [
        'Course source: middleEar.pptx, slides 16–20.',
        'J. Zwislocki, “Analysis of the Middle-Ear Function. Part I: Input Impedance,” JASA 34, 1514–1523 (1962), DOI: 10.1121/1.1918382.',
        'K. N. O’Connor and S. Puria, “Middle-ear circuit model parameters based on a population of human ears,” JASA 123, 197–211 (2008), DOI: 10.1121/1.2817358.',
        'E. Merchant and S. T. Neely, “Effects of middle ear pathology on middle ear impedance in chinchillas,” JASA 150, 969 (2021), DOI: 10.1121/10.0005822.',
        'E. Merchant and S. T. Neely, “A model of middle-ear impedance in children with otitis media with effusion,” Ear and Hearing (2022), DOI: 10.1097/AUD.0000000000001317.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(source)
        set_run_font(r, size=9.2, color=INK)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build_report()
